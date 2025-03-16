import zipfile
import xmltodict
import requests
import os
import shutil
from zipfile import ZipFile
from xml.dom.minidom import parse
import argparse
import torch
import torch.optim as optim
import torch.nn.functional as F
from fastNLP import Trainer
#wps
def wps_train(filename):
    transF = os.path.splitext(filename)
    new_name = 'test.txt'
    os.rename(filename, new_name)

#ofd
def ofd_train(filename):
    transF = os.path.splitext(filename)
    new_name = transF[0] + '.zip'
    os.rename(filename, new_name)
    path = '测试.zip'
    file = ZipFile(path)
    file.extractall()
    file.close()
    fp = open("test.txt", "w")
    dom = parse('Content.xml')
    path = ".\Doc_0\Pages"
    for file_name in os.listdir(path):
        dom = parse(path+'.\\'+file_name+r'\Content.xml')
        books = dom.getElementsByTagName('ofd:TextCode')
        for i in books:
            fp.write(i.firstChild.data)
            if (i.firstChild.data == '。'):
                fp.write('\n')

#word
def word(filename):
    import docx
    file=docx.Document(filename)
    fp = open("test.txt","w",encoding='utf-8')
    for i in range(len(file.paragraphs)):
        for j in file.paragraphs[i].text:
            fp.write(j)
            if j == '。':
                fp.write('\n')

#pdf



def pdf_docx(file_path):
    from pdf2docx import Converter
    files = []
    files.append(file_path)
    for file in files:
        suff_name = os.path.splitext(file)[1]

        if suff_name != '.pdf':
            continue
        file_name = os.path.splitext(file)[0]
        pdf_name = file
        docx_name = file_name + '.docx'
        cv = Converter(pdf_name)
        cv.convert(docx_name)
        cv.close()


def txt_train(filename):
    new_name = 'test.txt'
    os.rename(filename, new_name)